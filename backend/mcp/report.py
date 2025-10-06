from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
import boto3
import os
import uuid
from dotenv import load_dotenv

load_dotenv()

class ReportGenerator:
    def __init__(self):
        self.s3_client = boto3.client(
            's3',
            region_name=os.getenv('AWS_REGION', 'us-east-1'),
            aws_access_key_id=os.getenv('AWS_ACCESS_KEY_ID'),
            aws_secret_access_key=os.getenv('AWS_SECRET_ACCESS_KEY')
        )
        self.bucket = os.getenv('S3_BUCKET_REPORTS', 'mavik-reports')

    async def generate_docx(
        self,
        content: str = None,
        sections: List[Dict[str, Any]] = None,
        title: str = "Pre-Screening Analysis"
    ) -> str:
        """Generate Word document with 1\" margins and proper formatting

        Args:
            content: Markdown text to convert to Word (new format)
            sections: List of section dicts (legacy format)
            title: Document title
        """

        doc = Document()

        # Set 1" margins
        sections_obj = doc.sections
        for section in sections_obj:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Add title
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # New format: Convert markdown content to Word
        if content:
            self._add_markdown_to_doc(doc, content)

        # Legacy format: Add structured sections
        elif sections:
            for section in sections:
                # Section title
                doc.add_heading(section['title'], level=1)

                # Section content
                content_para = doc.add_paragraph(section['content'])
                content_para.paragraph_format.space_after = Pt(12)

        # Save to temp file (Windows-compatible)
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_filename = os.path.join(temp_dir, f"{uuid.uuid4()}.docx")
        doc.save(temp_filename)

        # Upload to S3
        s3_key = f"reports/{uuid.uuid4()}.docx"
        self.s3_client.upload_file(temp_filename, self.bucket, s3_key)

        # Generate presigned URL (valid for 7 days)
        url = self.s3_client.generate_presigned_url(
            'get_object',
            Params={'Bucket': self.bucket, 'Key': s3_key},
            ExpiresIn=604800  # 7 days
        )

        # Clean up temp file
        os.remove(temp_filename)

        return url

    def _add_markdown_to_doc(self, doc: Document, markdown_text: str):
        """Convert markdown text to Word document with proper formatting"""
        import re

        lines = markdown_text.split('\n')
        i = 0

        while i < len(lines):
            line = lines[i].strip()

            if not line:
                i += 1
                continue

            # H1 header (## Header)
            if line.startswith('## '):
                doc.add_heading(line[3:], level=1)

            # H2 header (### Header)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=2)

            # H3 header (#### Header)
            elif line.startswith('#### '):
                doc.add_heading(line[5:], level=3)

            # Bullet list (- item or * item)
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')

            # Numbered list (1. item)
            elif re.match(r'^\d+\.\s', line):
                content = re.sub(r'^\d+\.\s', '', line)
                doc.add_paragraph(content, style='List Number')

            # Bold text or regular paragraph
            else:
                para = doc.add_paragraph()

                # Handle inline formatting (bold)
                parts = re.split(r'(\*\*.*?\*\*)', line)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        # Bold text
                        para.add_run(part[2:-2]).bold = True
                    elif part:
                        # Regular text
                        para.add_run(part)

                para.paragraph_format.space_after = Pt(6)

            i += 1

# Global instance
report_generator = ReportGenerator()

async def generate_docx(content: str = None, sections: List[Dict[str, Any]] = None, title: str = "Pre-Screening Analysis") -> str:
    return await report_generator.generate_docx(content=content, sections=sections, title=title)
